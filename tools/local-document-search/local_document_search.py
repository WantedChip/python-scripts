"""Local Document Search Tool.

Privacy-first full-text search for your local files.
Indexes documents locally and provides fast search without uploading anything.
"""

# pylint: disable=too-many-lines,line-too-long,too-many-locals,too-many-arguments
# pylint: disable=too-many-positional-arguments,import-outside-toplevel

import argparse
import fnmatch
import hashlib
import json
import logging
import os
import re
import sqlite3
import sys
import time
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Dict, Generator, List, Set, Tuple

logger = logging.getLogger("local_document_search")


# -----------------------------------------------------------------------------
# Data Classes
# -----------------------------------------------------------------------------
@dataclass
class Document:
    """Represents a document in the index."""

    path: str
    relative_path: str
    extension: str
    size: int
    modified_time: float
    content_hash: str


@dataclass
class SearchResult:
    """Represents a search result."""

    document: Document
    score: float
    matched_terms: List[str]
    snippets: List[str]


# -----------------------------------------------------------------------------
# Text Extraction
# -----------------------------------------------------------------------------
def extract_text_from_file(filepath: Path) -> Tuple[str, str]:
    """Extract text content from a file based on its extension.

    Args:
        filepath: Path to the file.

    Returns:
        Tuple of (extracted_text, content_hash).
    """
    suffix = filepath.suffix.lower()

    # Text-based files - use stdlib
    text_extensions = {
        ".txt",
        ".md",
        ".markdown",
        ".rst",
        ".py",
        ".js",
        ".ts",
        ".jsx",
        ".tsx",
        ".html",
        ".htm",
        ".css",
        ".scss",
        ".sass",
        ".json",
        ".yaml",
        ".yml",
        ".toml",
        ".ini",
        ".cfg",
        ".conf",
        ".xml",
        ".csv",
        ".log",
        ".sql",
        ".sh",
        ".bash",
        ".zsh",
        ".fish",
        ".ps1",
        ".bat",
        ".cmd",
        ".dockerfile",
        ".gitignore",
        ".gitattributes",
        ".env",
        ".editorconfig",
        ".prettierrc",
        ".eslintrc",
        ".pylintrc",
        ".mypy.ini",
        ".flake8",
        ".isort.cfg",
    }

    if suffix in text_extensions or (
        suffix == ""
        and filepath.name
        in {"Dockerfile", "Makefile", "README", "LICENSE", "CHANGELOG"}
    ):
        return _extract_text_plain(filepath)

    # PDF - optional dependency
    if suffix == ".pdf":
        return _extract_text_pdf(filepath)

    # DOCX - optional dependency
    if suffix == ".docx":
        return _extract_text_docx(filepath)

    # Unsupported - return empty
    logger.debug("Skipping unsupported file type: %s", filepath)
    return "", ""


def _extract_text_plain(filepath: Path) -> Tuple[str, str]:
    """Extract text from plain text files with encoding detection."""
    # Try UTF-8 first, then latin-1
    for encoding in ("utf-8", "latin-1"):
        try:
            content = filepath.read_text(encoding=encoding)
            content_hash = hashlib.sha256(content.encode("utf-8")).hexdigest()[:16]
            return content, content_hash
        except UnicodeDecodeError:
            continue

    # Fallback: read as binary, decode with replace
    content = filepath.read_bytes().decode("utf-8", errors="replace")
    content_hash = hashlib.sha256(content.encode("utf-8")).hexdigest()[:16]
    return content, content_hash


def _extract_text_pdf(filepath: Path) -> Tuple[str, str]:
    """Extract text from PDF files using pypdf if available."""

    try:
        from pypdf import PdfReader
    except ImportError:
        logger.debug("pypdf not installed, skipping PDF: %s", filepath)
        return "", ""

    try:
        reader = PdfReader(str(filepath))
        text_parts = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
        content = "\n".join(text_parts)
        content_hash = hashlib.sha256(content.encode("utf-8")).hexdigest()[:16]
        return content, content_hash
    except Exception as e:  # pylint: disable=broad-except
        logger.warning("Failed to extract text from PDF %s: %s", filepath, e)
        return "", ""


def _extract_text_docx(filepath: Path) -> Tuple[str, str]:
    """Extract text from DOCX files using python-docx if available."""

    try:
        from docx import Document as DocxDocument
    except ImportError:
        logger.debug("python-docx not installed, skipping DOCX: %s", filepath)
        return "", ""

    try:
        doc = DocxDocument(str(filepath))
        text_parts = [para.text for para in doc.paragraphs if para.text.strip()]
        content = "\n".join(text_parts)
        content_hash = hashlib.sha256(content.encode("utf-8")).hexdigest()[:16]
        return content, content_hash
    except Exception as e:  # pylint: disable=broad-except
        logger.warning("Failed to extract text from DOCX %s: %s", filepath, e)
        return "", ""


# -----------------------------------------------------------------------------
# Text Processing
# -----------------------------------------------------------------------------
def tokenize(text: str) -> List[str]:
    """Tokenize text into searchable terms.

    Args:
        text: Input text.

    Returns:
        List of lowercase tokens (alphanumeric, 2+ chars).
    """
    # Split on non-alphanumeric, keep tokens 2+ chars
    tokens = re.findall(r"[a-zA-Z0-9]{2,}", text.lower())
    return tokens


def build_snippets(
    text: str, query_terms: List[str], max_snippets: int = 3, context_chars: int = 100
) -> List[str]:
    """Build context snippets around matched terms.

    Args:
        text: Full document text.
        query_terms: List of query terms to highlight.
        max_snippets: Maximum number of snippets to return.
        context_chars: Characters of context around each match.

    Returns:
        List of snippet strings.
    """
    if not text or not query_terms:
        return []

    snippets: List[str] = []
    text_lower = text.lower()

    for term in query_terms:
        term_lower = term.lower()
        start = 0
        while len(snippets) < max_snippets:
            idx = text_lower.find(term_lower, start)
            if idx == -1:
                break
            snippet_start = max(0, idx - context_chars)
            snippet_end = min(len(text), idx + len(term) + context_chars)
            snippet = text[snippet_start:snippet_end]
            if snippet_start > 0:
                snippet = "..." + snippet
            if snippet_end < len(text):
                snippet = snippet + "..."
            snippets.append(snippet)
            start = idx + 1

    return snippets[:max_snippets]


# -----------------------------------------------------------------------------
# Indexing
# -----------------------------------------------------------------------------
class DocumentIndex:
    """Inverted index for document search using SQLite."""

    def __init__(self, db_path: Path):
        """Initialize the index.

        Args:
            db_path: Path to SQLite database file.
        """
        self.db_path = db_path
        self._init_db()

    def _init_db(self) -> None:
        """Create database schema if not exists."""
        with sqlite3.connect(self.db_path) as conn:
            conn.execute(
                """
                CREATE TABLE IF NOT EXISTS documents (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    path TEXT UNIQUE NOT NULL,
                    relative_path TEXT NOT NULL,
                    extension TEXT NOT NULL,
                    size INTEGER NOT NULL,
                    modified_time REAL NOT NULL,
                    content_hash TEXT NOT NULL,
                    indexed_time REAL NOT NULL
                )
            """
            )
            conn.execute(
                """
                CREATE TABLE IF NOT EXISTS terms (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    term TEXT UNIQUE NOT NULL
                )
            """
            )
            conn.execute(
                """
                CREATE TABLE IF NOT EXISTS postings (
                    term_id INTEGER NOT NULL,
                    doc_id INTEGER NOT NULL,
                    frequency INTEGER NOT NULL DEFAULT 1,
                    positions TEXT,  -- JSON array of positions
                    PRIMARY KEY (term_id, doc_id),
                    FOREIGN KEY (term_id) REFERENCES terms(id),
                    FOREIGN KEY (doc_id) REFERENCES documents(id)
                )
            """
            )
            conn.execute(
                "CREATE INDEX IF NOT EXISTS idx_documents_path ON documents(path)"
            )
            conn.execute("CREATE INDEX IF NOT EXISTS idx_terms_term ON terms(term)")
            conn.execute(
                "CREATE INDEX IF NOT EXISTS idx_postings_term ON postings(term_id)"
            )
            conn.execute(
                "CREATE INDEX IF NOT EXISTS idx_postings_doc ON postings(doc_id)"
            )

    def add_document(self, doc: Document, tokens: List[str]) -> int:
        """Add or update a document in the index.

        Args:
            doc: Document metadata.
            tokens: List of tokens from document content.

        Returns:
            Document ID.
        """
        with sqlite3.connect(self.db_path) as conn:
            # Check if document exists and content unchanged
            cursor = conn.execute(
                "SELECT id, content_hash FROM documents WHERE path = ?", (doc.path,)
            )
            row = cursor.fetchone()

            if row and row[1] == doc.content_hash:
                logger.debug("Document unchanged, skipping: %s", doc.path)
                return int(row[0])

            # Insert or update document
            if row:
                doc_id = row[0]
                conn.execute(
                    """
                    UPDATE documents
                    SET relative_path=?, extension=?, size=?, modified_time=?,
                        content_hash=?, indexed_time=? WHERE id=?
                    """,
                    (
                        doc.relative_path,
                        doc.extension,
                        doc.size,
                        doc.modified_time,
                        doc.content_hash,
                        time.time(),
                        doc_id,
                    ),
                )
                # Remove old postings
                conn.execute("DELETE FROM postings WHERE doc_id = ?", (doc_id,))
            else:
                cursor = conn.execute(
                    """
                    INSERT INTO documents
                    (path, relative_path, extension, size, modified_time,
                     content_hash, indexed_time)
                    VALUES (?, ?, ?, ?, ?, ?, ?)
                    """,
                    (
                        doc.path,
                        doc.relative_path,
                        doc.extension,
                        doc.size,
                        doc.modified_time,
                        doc.content_hash,
                        time.time(),
                    ),
                )
                doc_id = cursor.lastrowid

            # Add term postings
            term_counts: Dict[str, int] = {}
            term_positions: Dict[str, List[int]] = {}

            for pos, token in enumerate(tokens):
                term_counts[token] = term_counts.get(token, 0) + 1
                if token not in term_positions:
                    term_positions[token] = []
                term_positions[token].append(pos)

            for term, freq in term_counts.items():
                # Get or create term
                cursor = conn.execute("SELECT id FROM terms WHERE term = ?", (term,))
                term_row = cursor.fetchone()
                if term_row:
                    term_id = term_row[0]
                else:
                    cursor = conn.execute(
                        "INSERT INTO terms (term) VALUES (?)", (term,)
                    )
                    term_id = cursor.lastrowid

                # Insert posting
                positions_json = json.dumps(term_positions[term])
                conn.execute(
                    "INSERT INTO postings (term_id, doc_id, frequency, positions)"
                    " VALUES (?, ?, ?, ?)",
                    (term_id, doc_id, freq, positions_json),
                )

            conn.commit()
            return int(doc_id)

    def remove_document(self, path: str) -> bool:
        """Remove a document from the index.

        Args:
            path: Absolute path of document to remove.

        Returns:
            True if document was found and removed.
        """
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.execute("SELECT id FROM documents WHERE path = ?", (path,))
            row = cursor.fetchone()
            if not row:
                return False
            doc_id = row[0]
            conn.execute("DELETE FROM postings WHERE doc_id = ?", (doc_id,))
            conn.execute("DELETE FROM documents WHERE id = ?", (doc_id,))
            conn.commit()
            return True

    def search(self, query: str, limit: int = 20) -> List[SearchResult]:
        """Search for documents matching query.

        Args:
            query: Search query string.
            limit: Maximum number of results.

        Returns:
            List of search results sorted by relevance.
        """
        query_tokens = tokenize(query)
        if not query_tokens:
            return []

        with sqlite3.connect(self.db_path) as conn:
            conn.row_factory = sqlite3.Row

            # Get term IDs
            placeholders = ",".join("?" * len(query_tokens))
            # fmt: off
            sql = (
                f"SELECT id, term FROM terms WHERE term IN"
                f" ({placeholders})"  # nosec B608
            )
            # fmt: on
            cursor = conn.execute(sql, query_tokens)
            term_rows = cursor.fetchall()
            if not term_rows:
                return []

            term_map = {row["term"]: row["id"] for row in term_rows}

            # Get documents matching any query term with frequency
            term_ids = list(term_map.values())
            placeholders = ",".join("?" * len(term_ids))
            sql = (
                f"SELECT p.doc_id, d.path, d.relative_path, d.extension, d.size,"
                f" d.modified_time, d.content_hash, p.term_id, p.frequency,"
                f" p.positions, t.term"
                f" FROM postings p"
                f" JOIN documents d ON p.doc_id = d.id"
                f" JOIN terms t ON p.term_id = t.id"
                f" WHERE p.term_id IN ({placeholders})"  # nosec B608
            )
            cursor = conn.execute(sql, term_ids)
            rows = cursor.fetchall()

            # Aggregate scores per document
            doc_scores: Dict[int, Dict[str, Any]] = {}
            for row in rows:
                doc_id = row["doc_id"]
                if doc_id not in doc_scores:
                    doc_scores[doc_id] = {
                        "doc": Document(
                            path=row["path"],
                            relative_path=row["relative_path"],
                            extension=row["extension"],
                            size=row["size"],
                            modified_time=row["modified_time"],
                            content_hash=row["content_hash"],
                        ),
                        "score": 0.0,
                        "matched_terms": set(),
                        "all_positions": [],
                    }
                doc_scores[doc_id]["score"] += row["frequency"]
                doc_scores[doc_id]["matched_terms"].add(row["term"])
                positions = json.loads(row["positions"]) if row["positions"] else []
                doc_scores[doc_id]["all_positions"].extend(positions)

            # Sort by score descending
            sorted_docs = sorted(
                doc_scores.values(),
                key=lambda x: x["score"],
                reverse=True,
            )[:limit]

            # Build results with snippets
            results = []
            for entry in sorted_docs:
                # Get full text for snippets
                full_text = self._get_document_text(entry["doc"].path)
                snippets = build_snippets(full_text, list(entry["matched_terms"]))
                results.append(
                    SearchResult(
                        document=entry["doc"],
                        score=entry["score"],
                        matched_terms=list(entry["matched_terms"]),
                        snippets=snippets,
                    )
                )

            return results

    def _get_document_text(self, path: str) -> str:
        """Get full text of a document by path."""
        try:
            content, _ = extract_text_from_file(Path(path))
            return content
        except Exception:  # pylint: disable=broad-except
            return ""

    def get_stats(self) -> Dict[str, Any]:
        """Get index statistics."""
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.execute("SELECT COUNT(*) FROM documents")
            doc_count = cursor.fetchone()[0]
            cursor = conn.execute("SELECT COUNT(*) FROM terms")
            term_count = cursor.fetchone()[0]
            cursor = conn.execute("SELECT COUNT(*) FROM postings")
            posting_count = cursor.fetchone()[0]
            return {
                "documents": doc_count,
                "unique_terms": term_count,
                "postings": posting_count,
                "db_size_mb": (
                    self.db_path.stat().st_size / (1024 * 1024)
                    if self.db_path.exists()
                    else 0
                ),
            }

    def get_all_document_paths(self) -> Set[str]:
        """Get all indexed document paths."""
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.execute("SELECT path FROM documents")
            return {row[0] for row in cursor.fetchall()}


# -----------------------------------------------------------------------------
# File Discovery
# -----------------------------------------------------------------------------
def should_index_file(
    filepath: Path,
    include_patterns: List[str],
    exclude_patterns: List[str],
    max_size_mb: float,
) -> bool:
    """Check if a file should be indexed.

    Args:
        filepath: Path to file.
        include_patterns: Glob patterns to include.
        exclude_patterns: Glob patterns to exclude.
        max_size_mb: Maximum file size in MB.

    Returns:
        True if file should be indexed.
    """
    # Check size
    try:
        size_mb = filepath.stat().st_size / (1024 * 1024)
        if size_mb > max_size_mb:
            return False
    except OSError:
        return False

    # Check include patterns
    if include_patterns:
        matched = False
        for pattern in include_patterns:
            if fnmatch.fnmatch(filepath.name, pattern) or fnmatch.fnmatch(
                str(filepath), pattern
            ):
                matched = True
                break
        if not matched:
            return False

    # Check exclude patterns
    for pattern in exclude_patterns:
        if fnmatch.fnmatch(filepath.name, pattern) or fnmatch.fnmatch(
            str(filepath), pattern
        ):
            return False

    return True


def discover_files(
    root: Path,
    include_patterns: List[str],
    exclude_patterns: List[str],
    max_size_mb: float,
    follow_symlinks: bool = False,
) -> Generator[Path, None, None]:
    """Discover files to index.

    Args:
        root: Root directory to search.
        include_patterns: Glob patterns to include.
        exclude_patterns: Glob patterns to exclude.
        max_size_mb: Maximum file size in MB.
        follow_symlinks: Whether to follow symlinks.

    Yields:
        Paths of files to index.
    """
    exclude_dirs = {
        ".git",
        "__pycache__",
        ".venv",
        "venv",
        "node_modules",
        ".pytest_cache",
        ".mypy_cache",
        ".ruff_cache",
        "dist",
        "build",
        ".idea",
        ".vscode",
        ".vs",
        "target",
        "bin",
        "obj",
    }

    for dirpath, dirnames, filenames in os.walk(root, followlinks=follow_symlinks):
        # Filter out excluded directories
        dirnames[:] = [d for d in dirnames if d not in exclude_dirs]

        for filename in filenames:
            filepath = Path(dirpath) / filename
            if should_index_file(
                filepath, include_patterns, exclude_patterns, max_size_mb
            ):
                yield filepath


# -----------------------------------------------------------------------------
# Indexing Pipeline
# -----------------------------------------------------------------------------
def index_directory(
    root: Path,
    index: DocumentIndex,
    include_patterns: List[str],
    exclude_patterns: List[str],
    max_size_mb: float,
    force_reindex: bool = False,
    verbose: bool = False,
) -> Tuple[int, int]:
    """Index all files in a directory.

    Args:
        root: Root directory to index.
        index: DocumentIndex instance.
        include_patterns: Glob patterns to include.
        exclude_patterns: Glob patterns to exclude.
        max_size_mb: Maximum file size in MB.
        force_reindex: Reindex even if unchanged.
        verbose: Verbose logging.

    Returns:
        Tuple of (indexed_count, skipped_count).
    """
    indexed = 0
    skipped = 0
    errors = 0

    # Get already indexed paths for incremental indexing
    indexed_paths = index.get_all_document_paths() if not force_reindex else set()

    for filepath in discover_files(
        root, include_patterns, exclude_patterns, max_size_mb
    ):
        abs_path = str(filepath.resolve())

        if not force_reindex and abs_path in indexed_paths:
            # Check if modified
            try:
                stat = filepath.stat()
                with sqlite3.connect(index.db_path) as conn:
                    cursor = conn.execute(
                        "SELECT modified_time, content_hash FROM documents"
                        " WHERE path = ?",
                        (abs_path,),
                    )
                    row = cursor.fetchone()
                    if row and row[0] == stat.st_mtime:
                        skipped += 1
                        continue
            except OSError:
                pass

        try:
            content, content_hash = extract_text_from_file(filepath)
            if not content.strip():
                skipped += 1
                continue

            tokens = tokenize(content)
            if not tokens:
                skipped += 1
                continue

            stat = filepath.stat()
            rel_path = filepath.relative_to(root)

            doc = Document(
                path=abs_path,
                relative_path=str(rel_path),
                extension=filepath.suffix.lower(),
                size=stat.st_size,
                modified_time=stat.st_mtime,
                content_hash=content_hash,
            )

            index.add_document(doc, tokens)
            indexed += 1

            if verbose:
                logger.info("Indexed: %s (%d tokens)", rel_path, len(tokens))

        except Exception as e:  # pylint: disable=broad-except
            errors += 1
            logger.warning("Failed to index %s: %s", filepath, e)

    if errors:
        logger.warning("Completed with %d errors", errors)

    return indexed, skipped


# -----------------------------------------------------------------------------
# CLI
# -----------------------------------------------------------------------------
def setup_logging(verbose: bool) -> None:
    """Configure logging."""
    level = logging.DEBUG if verbose else logging.INFO
    handler = logging.StreamHandler(sys.stderr)
    handler.setFormatter(
        logging.Formatter(
            "%(asctime)s - %(levelname)s - %(message)s", datefmt="%Y-%m-%d %H:%M:%S"
        )
    )
    logger.setLevel(level)
    logger.addHandler(handler)
    logger.propagate = False


def cmd_index(args: argparse.Namespace) -> int:
    """Index command handler."""
    root = Path(args.path).resolve()
    if not root.exists():
        logger.error("Path does not exist: %s", root)
        return 1
    if not root.is_dir():
        logger.error("Path is not a directory: %s", root)
        return 1

    index_path = Path(args.index).resolve()
    index_path.parent.mkdir(parents=True, exist_ok=True)

    index = DocumentIndex(index_path)

    include = args.include.split(",") if args.include else []
    exclude = args.exclude.split(",") if args.exclude else []

    logger.info("Indexing directory: %s", root)
    logger.info("Index location: %s", index_path)
    if include:
        logger.info("Include patterns: %s", include)
    if exclude:
        logger.info("Exclude patterns: %s", exclude)

    start = time.time()
    indexed, skipped = index_directory(
        root,
        index,
        include,
        exclude,
        args.max_size,
        force_reindex=args.force,
        verbose=args.verbose,
    )
    elapsed = time.time() - start

    stats = index.get_stats()
    logger.info("Indexing complete in %.2fs", elapsed)
    logger.info("Documents indexed: %d, skipped: %d", indexed, skipped)
    logger.info(
        "Index stats: %d docs, %d terms, %.2f MB",
        stats["documents"],
        stats["unique_terms"],
        stats["db_size_mb"],
    )

    return 0


def cmd_search(args: argparse.Namespace) -> int:
    """Search command handler."""
    index_path = Path(args.index).resolve()
    if not index_path.exists():
        logger.error("Index not found: %s. Run 'index' first.", index_path)
        return 1

    index = DocumentIndex(index_path)

    results = index.search(args.query, limit=args.limit)

    if not results:
        logger.info("No results found for: %s", args.query)
        return 0

    # Output format
    if args.format == "json":
        output = []
        for r in results:
            output.append(
                {
                    "path": r.document.path,
                    "relative_path": r.document.relative_path,
                    "extension": r.document.extension,
                    "size": r.document.size,
                    "score": r.score,
                    "matched_terms": r.matched_terms,
                    "snippets": r.snippets,
                }
            )
        print(json.dumps(output, indent=2))
    else:
        print(f"Found {len(results)} result(s) for: {args.query}\n")
        for i, r in enumerate(results, 1):
            print(f"{i}. {r.document.relative_path} (score: {r.score:.1f})")
            print(f"   Path: {r.document.path}")
            print(f"   Size: {r.document.size:,} bytes | Ext: {r.document.extension}")
            print(f"   Matched: {', '.join(r.matched_terms)}")
            if r.snippets:
                for snippet in r.snippets:
                    print(f"   ... {snippet}")
            print()

    return 0


def cmd_stats(args: argparse.Namespace) -> int:
    """Stats command handler."""
    index_path = Path(args.index).resolve()
    if not index_path.exists():
        logger.error("Index not found: %s", index_path)
        return 1

    index = DocumentIndex(index_path)
    stats = index.get_stats()

    if args.format == "json":
        print(json.dumps(stats, indent=2))
    else:
        print("Index Statistics:")
        print(f"  Documents: {stats['documents']}")
        print(f"  Unique terms: {stats['unique_terms']}")
        print(f"  Postings: {stats['postings']}")
        print(f"  Database size: {stats['db_size_mb']:.2f} MB")
        print(f"  Index file: {index_path}")

    return 0


def cmd_remove(args: argparse.Namespace) -> int:
    """Remove command handler."""
    index_path = Path(args.index).resolve()
    if not index_path.exists():
        logger.error("Index not found: %s", index_path)
        return 1

    index = DocumentIndex(index_path)
    removed = index.remove_document(args.path)

    if removed:
        logger.info("Removed from index: %s", args.path)
    else:
        logger.warning("Document not found in index: %s", args.path)

    return 0


def main() -> int:
    """Main entry point."""
    parser = argparse.ArgumentParser(
        description=(
            "Privacy-first local document search - "
            "index and search your files locally."
        ),
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Index a directory
  python local_document_search.py index ~/Documents
      --index ~/.local/share/docsearch/index.db

  # Search for terms
  python local_document_search.py search "machine learning"
      --index ~/.local/share/docsearch/index.db

  # Search with JSON output
  python local_document_search.py search "config"
      --format json --limit 5

  # Show index statistics
  python local_document_search.py stats
      --index ~/.local/share/docsearch/index.db

  # Reindex (force)
  python local_document_search.py index ~/Documents --force

Notes:
  - All processing happens locally. No data leaves your machine.
        """,
    )

    parser.add_argument(
        "-v", "--verbose", action="store_true", help="Enable verbose logging"
    )

    subparsers = parser.add_subparsers(dest="command", required=True)

    # Index command
    idx_parser = subparsers.add_parser("index", help="Index a directory")
    idx_parser.add_argument("path", help="Directory to index")
    idx_parser.add_argument(
        "--include", help="Comma-separated glob patterns to include"
    )
    idx_parser.add_argument(
        "--exclude", help="Comma-separated glob patterns to exclude"
    )
    idx_parser.add_argument(
        "--max-size",
        type=float,
        default=50,
        help="Maximum file size in MB (default: 50)",
    )
    idx_parser.add_argument(
        "--force", action="store_true", help="Force reindex all files"
    )
    idx_parser.add_argument(
        "--index",
        default="~/.local/share/docsearch/index.db",
        help="Path to index database",
    )

    # Search command
    search_parser = subparsers.add_parser("search", help="Search the index")
    search_parser.add_argument("query", help="Search query")
    search_parser.add_argument(
        "-l", "--limit", type=int, default=20, help="Max results (default: 20)"
    )
    search_parser.add_argument(
        "--format", choices=["text", "json"], default="text", help="Output format"
    )
    search_parser.add_argument(
        "--index",
        default="~/.local/share/docsearch/index.db",
        help="Path to index database",
    )

    # Stats command
    stats_parser = subparsers.add_parser("stats", help="Show index statistics")
    stats_parser.add_argument("--format", choices=["text", "json"], default="text")
    stats_parser.add_argument(
        "--index",
        default="~/.local/share/docsearch/index.db",
        help="Path to index database",
    )

    # Remove command
    remove_parser = subparsers.add_parser("remove", help="Remove a document from index")
    remove_parser.add_argument("path", help="Path of document to remove")
    remove_parser.add_argument(
        "--index",
        default="~/.local/share/docsearch/index.db",
        help="Path to index database",
    )

    args = parser.parse_args()
    setup_logging(args.verbose)

    # Expand index path
    args.index = os.path.expanduser(args.index)

    if args.command == "index":
        return cmd_index(args)
    if args.command == "search":
        return cmd_search(args)
    if args.command == "stats":
        return cmd_stats(args)
    if args.command == "remove":
        return cmd_remove(args)

    return 1


if __name__ == "__main__":
    sys.exit(main())
